import gettext
import locale
from typing import Callable, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
import io
from pathlib import Path


__all__ = ["ReportGenerator"]

class LocalizationHelper:
    """
    Handles translation and localization using gettext.

    Parameters
    ----------
    language : str, optional
        The language code to use for localization (default is 'en').
    locale_dir : str, optional
        The directory where localization files are stored (default is 
        "pump/utilites/locales"). If changed, insert full path.
    """
    def __init__(self, language: str = 'en', locale_dir: Optional[str] = None) -> None:
        self.language = language
        self.locale_dir = locale_dir
        self._ = self._setup_translation()

    def _setup_translation(self) -> Callable[[str], str]:
        """
        Sets up gettext for localization.

        Returns
        -------
        Callable[[str], str]
            A function that translates a given string based on the selected language.
        """
        try:
            locale.setlocale(locale.LC_ALL, '')
        except locale.Error:
            print("Warning: Unable to set locale, using default.")

        try:
            if self.locale_dir:
                locale_dir = self.locale_dir
            else:
                locale_dir = Path(__file__).parent / "locales"
            lang_translation = gettext.translation(
                "messages",
                localedir=locale_dir,
                languages=[self.language],
                fallback=True
            )
            lang_translation.install()
            return lang_translation.gettext
        except FileNotFoundError:
            return gettext.gettext  # Fallback to default gettext
        

class ReportGenerator:
    """
    Generates a report from predefined templates using the `python-docx` library.

    This class loads a document template, populates it with data, and applies localization
    to ensure language-specific content where necessary.

    Parameters
    ----------
    language : str, optional
        The language for the report (default is 'en').
    locale_dir : str, optional
        The directory containing localization files (default is 
        "pump/utilites/locales"). If changed, insert full path.
    template_dir : str, optional
        The directory containing report templates (default is "templates").

    Attributes
    ----------
    language : str
        Selected language code.
    template_dir : str
        Directory containing document templates.
    localization : LocalizationHelper
        Handles language-specific text translations.
    _ : Callable[[str], str]
        Translation function from the localization helper.
    """

    def __init__(self, language='en', locale_dir : Optional[str]=None, template_path : Optional[str]=None):
        self.language = language
        self.template_path = template_path
        self.localization = LocalizationHelper(language, locale_dir)
        self._ = self.localization._

    def _get_template_path(self):
        """
        Retrieve the appropriate template file based on the selected language.
        
        Returns
        -------
        str
            The full path to the template file.
        
        Raises
        ------
        FileNotFoundError
            If the specified template file is not found.
        """
        if self.template_path:
            if not os.path.exists(self.template_path):
                raise FileNotFoundError(f"Template file not found: {self.template_path}")
            return self.template_path
        else:
            template_file = f"template_{self.language}.docx"
            template_path = Path(__file__).parents[1] / "templates" / template_file
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template file not found: {template_path}")
            return template_path

    def generate_report(self, report_data, output_file=None):
        """
        Generates a report from a predefined template.

        Parameters
        ----------
        report_data : dict
            Dictionary containing all necessary report information.
        output_file : str, optional
            Path to save the generated report.

        Returns
        -------
        str
            Path of the saved report file.
        """
        TAGS = report_data["equipment_description"]["TAG"]
        if output_file is None:
            if "/" in TAGS:
                str_TAGS = TAGS.replace("/", "")
            else:
                str_TAGS = TAGS
            output_file = f"{self._("Report")}_{str_TAGS}_{datetime.now().strftime('%d-%m-%Y')}.docx"

        template_path = self._get_template_path()
        document = Document(template_path)

        self._add_intro(document)
        self._add_equipment_description(document, report_data.get("equipment_description", {}))
        self._add_design_point(document, report_data.get("design_point", {}))
        for tag, test_data in report_data["test_data"].items():
            document.add_paragraph(self._('Test Data - ') + tag, style='Title_2')
            self._add_test_summary(document, test_data.get("test_summary", {}))
            self._add_test_data(document, test_data.get("test_data", {}))

            # Adding  charts, if any.
            search_key = 'Curve'
            matches = [key for key in test_data.keys() if search_key in key]

            for chart in matches:
                self._add_chart(document, chart, test_data[chart])

        document.save(output_file)
        print(f"{self._('New document saved at')}: {output_file}")
        return output_file

    def _add_intro(self, document):
        """
        Add introductory section to the document.
        
        Parameters
        ----------
        document : docx.Document
            The Word document object to modify.
        """
        # Adding instructions to update the TOC
        TOC_note = (
            "-X-X-X-\n"
            "Note: To update the Table of Contents, open the document in Word, "
            "right-click on the TOC, and select 'Update Field'.\n"
            "-X-X-X-"
        )
        document.add_paragraph(self._(TOC_note), style='Text_body')
        
        # Note that the template already has an Title 1 called Introduction
        introduction_text_1 = (
            "This report documents the Factory Acceptance Test (FAT) performed on an API 610 "
            "centrifugal pump at the manufacturer's facility. The FAT was conducted to verify "
            "compliance with the technical specifications and contractual requirements before shipment."
        )
        document.add_paragraph(self._(introduction_text_1)).style = 'Text_body'
        introduction_text_2 = (
            "The primary objective of the FAT is to ensure that the pump meets the required  "
            "mechanical, and performance criteria as defined in API 610 and the project specifications. "
            "The test procedures include mechanical run tests and performance verification, vibration "
            "analysis, and NPSH (Net Positive Suction Head) testing."
        )
        document.add_paragraph(self._(introduction_text_2)).style = 'Text_body'
        introduction_text_3 = (
            "This report provides a detailed record of the test procedures, observed results, and "
            "compliance with the acceptance criteria. Any deviations, corrective actions, and final "
            "acceptance status are also documented to ensure transparency and traceability of the "
            "qualification process."
        )
        document.add_paragraph(self._(introduction_text_3)).style = 'Text_body'

        document.add_paragraph(self._("Tests Development"), style='Title_1')
        document.add_paragraph(
            self._("Here we present the key details regarding the development of the tests."), 
            style='Text_body'
        )

    def _add_design_point(self, document, design_point):
        """
        Add design point section to the document.
        
        Parameters
        ----------
        document : docx.Document
            The Word document object to modify.
        equipment_description : dict
            Dictionary containing equipment details.
        """
        table_font = "Consolas"
        width = Inches(2)
        if design_point:
            data = design_point.fluid.__dict__ | design_point.__dict__
            del data["fluid"]
            document.add_paragraph(self._(" "), style='Text_body')
            document.add_paragraph(self._('Design Point data:'), style='Text_body')
            table = document.add_table(rows=len(data), cols=2)
            table.style = "Table Grid"
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for row, (key, value) in enumerate(data.items()):
                cell_left = table.cell(row, 0)
                cell_rigth = table.cell(row, 1)
                cell_left.text = self._(key.replace("_", " ").title()) if key != "name" else "Service"
                cell_rigth.text = value if isinstance(value, str) else f"{value:0.02f~P}"
                cell_left.width = width
                cell_rigth.width = width
                run_left =  cell_left.paragraphs[0].runs[0]
                run_rigth =  cell_rigth.paragraphs[0].runs[0]
                run_left.bold = True  # Make title row boldface
                run_left.font.size = Pt(9)  # Set font size
                run_rigth.font.size = Pt(8)  # Set font size
                run_left.font.name = table_font  # Use fixed-spacing font
                run_rigth.font.name = table_font  # Use fixed-spacing font
        
    def _add_equipment_description(self, document, equipment_description):
        """
        Add equipment description section to the document.
        
        Parameters
        ----------
        document : docx.Document
            The Word document object to modify.
        equipment_description : dict
            Dictionary containing equipment details.
        """
        table_font = "Consolas"
        width = Inches(2)
        if equipment_description:
            document.add_paragraph(self._('Equipment Description'), style='Title_2')
            document.add_paragraph(self._("The following section presents the main characteristics of the equipment."), style='Text_body')
            table = document.add_table(rows=len(equipment_description), cols=2)
            table.style = "Table Grid"
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for row, (key, value) in enumerate(equipment_description.items()):
                cell_left = table.cell(row, 0)
                cell_rigth = table.cell(row, 1)
                cell_left.text = self._(key)
                cell_rigth.text = value
                cell_left.width = width
                cell_rigth.width = width
                run_left =  cell_left.paragraphs[0].runs[0]
                run_rigth =  cell_rigth.paragraphs[0].runs[0]
                run_left.bold = True  # Make title row boldface
                run_left.font.size = Pt(9)  # Set font size
                run_rigth.font.size = Pt(8)  # Set font size
                run_left.font.name = table_font  # Use fixed-spacing font
                run_rigth.font.name = table_font  # Use fixed-spacing font

    def _add_test_summary(self, document, test_summary):
        """
        Add test summary section to the document.
        
        Parameters
        ----------
        document : docx.Document
            The Word document object to modify.
        test_summary : dict
            Dictionary containing test results and observations.
        """
        table_font = "Consolas"
        width = Inches(1.25)
        if test_summary:
            document.add_paragraph(self._('Performance test summary.'), style='Title_3')
            test_description = (
                "The computations in this section were performed in accordancy with API 610 12th issue:"
            )
            document.add_paragraph(self._(test_description), style='Text_body')
            api_paragraph = (
                "8.3.3.4.3 The test data shall be fit to a spline or an appropriate polynomial (typically, "
                "not less than a third order) for head and for power using a least squares method. The "
                "resulting polynomial equation shall be stated on the head and power calculated. These "
                "values shall be corrected for speed, viscosity, and density (specific gravity).[..]"
            )
            document.add_paragraph(api_paragraph, style="Intense Quote")
            document.add_paragraph(self._("Calculated Head and Power:"), style='Text_body')
            table = document.add_table(rows=len(test_summary) + 1, cols=4)
            table.style = "Table Grid"
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Fill the first column with row names
            col_names = ["Parameter", "Actual", "Minimum", "Maximum"]
            
            
            for col, name in enumerate(col_names):
                cell = table.cell(0, col)
                cell.text = self._(name)
                cell.width = width
                run =  cell.paragraphs[0].runs[0]
                run.bold = True  # Make title row boldface
                run.font.size = Pt(9)  # Set font size
                run.font.name = table_font  # Use fixed-spacing font

                if col != 0:
                    table.cell(0, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for row, (parameter, value) in enumerate(test_summary.items(), start=1):
                cell_left = table.cell(row, 0)
                cell_left.text = self._(parameter)
                cell_left.width = width
                run =  cell_left.paragraphs[0].runs[0]
                # run.bold = True  # Make title row boldface
                run.font.size = Pt(8)  # Set font size
                run.font.name = table_font  # Use fixed-spacing font

                if isinstance(value, list):  # If value is a list
                    actual, min_val, max_val = value if len(value) == 3 else [value[0], "-", "-"]
                else:
                    actual, min_val, max_val = value, "-", "-"

                for col, text in enumerate([actual, min_val, max_val], start=1):
                    cell = table.cell(row, col)
                    cell.text = f"{text:0.02f~P}" if text != "-" else text
                    cell.width = width
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(8)  # Set font size
                    run.font.name = table_font  # Use fixed-spacing font
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_test_data(self, document, test_data):
        """
        Add test data section to the document.
        
        Parameters
        ----------
        document : docx.Document
            The Word document object to modify.
        test_data : dict
            Dictionary containing test results and observations.
        """
        table_font = "Consolas"
        if test_data:
            test_description = (
                "Here we can find the translation of the test results into data "
                "based on the specified speed of rotation (or frequency) and density."
            )
            document.add_paragraph(" ", style='Text_body')
            document.add_paragraph(self._(test_description), style='Text_body')           
            
            table = document.add_table(rows=len(test_data["Capacity"]) + 1, cols=len(test_data))
            table.style = "Table Grid"
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Fill the first column with row names
            col_names = list(test_data.keys())
            
            width = Inches(1.25)
            for col, name in enumerate(col_names):
                upper_cell = table.cell(0, col)
                upper_cell.text = self._(name)
                upper_cell.width = width
                run = upper_cell.paragraphs[0].runs[0]
                run.bold = True  # Make title row boldface
                run.font.size = Pt(9)  # Set font size
                run.font.name = table_font  # Use fixed-spacing font
                upper_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for row, value in enumerate(test_data[name], start=1):
                    cell = table.cell(row, col)
                    cell.text = f"{value:0.02f~P}"
                    cell.width = width
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(8)  # Set font size
                    run.font.name = table_font  # Use fixed-spacing font
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
                
    def _add_chart(self, document, chart_title, chart):
        """
        Add charts to the document.
        
        Parameters
        ----------
        document : docx.Document
            The Word document object to modify.
        charts : dict
            Dictionary containing charts in file paths or bytes format.
        """
        width = Inches(6)
        height = Inches(6)
        document.add_paragraph(self._(chart_title), style='Title_3')
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        if isinstance(chart, str) and os.path.exists(chart):  # File path
            run.add_picture(chart,
                            width=width,
                            height=height)
        elif isinstance(chart, io.BytesIO):  # In-memory buffer
            run.add_picture(chart,
                            width=width,
                            height=height)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

